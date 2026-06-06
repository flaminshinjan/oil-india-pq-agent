"""Document extractors for .docx, .xlsx, .pdf.

Each extractor returns a list of `Chunk` objects with text + metadata. The
chunks are deliberately self-describing — every chunk's text starts with
breadcrumbs (file name, section, table id, sheet, page) so the LLM sees the
context even after dense retrieval.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
import openpyxl
import pdfplumber


# ---------- data types ----------

@dataclass
class Chunk:
    text: str
    metadata: dict = field(default_factory=dict)


# ---------- helpers ----------

def _clean(s: str | None) -> str:
    if s is None:
        return ""
    s = str(s)
    s = s.replace(" ", " ").replace("\r", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _table_to_markdown(rows: list[list[str]]) -> str:
    """Render a list-of-rows table as a compact markdown table."""
    rows = [[_clean(c) for c in r] for r in rows if any((c or "").strip() for c in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []
    out = ["| " + " | ".join(header) + " |",
           "| " + " | ".join(["---"] * width) + " |"]
    for r in body:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


# ---------- .docx ----------

def extract_docx(path: Path) -> list[Chunk]:
    """Pull paragraphs + tables out of a .docx.

    PQ docx files use tables to store Q/A pairs. We render each table as
    markdown and emit one chunk per table, plus one chunk for the surrounding
    narrative paragraphs (header / signatory block / preamble).
    """
    doc = Document(str(path))
    chunks: list[Chunk] = []

    # Walk the body in document order so paragraphs that surround a table
    # land in the same chunk as the table when small enough.
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        t = _clean(para.text)
        if t:
            paragraphs.append(t)
    narrative = "\n".join(paragraphs)
    if narrative:
        chunks.append(Chunk(
            text=narrative,
            metadata={"section": "narrative"},
        ))

    for ti, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([_clean(cell.text) for cell in row.cells])
        md = _table_to_markdown(rows)
        if md:
            chunks.append(Chunk(
                text=f"Table {ti + 1}:\n{md}",
                metadata={"section": f"table_{ti + 1}"},
            ))

    return chunks


# ---------- .xlsx ----------

def extract_xlsx(path: Path) -> list[Chunk]:
    """One chunk per sheet. We render the whole sheet as a markdown table —
    these workbooks are small (tens of rows, <20 columns) so a full render
    preserves cross-row context the LLM needs for year-over-year questions.

    We also prepend a one-line natural-language descriptor of what columns
    the sheet contains, so dense retrieval has English to match against
    instead of just unit codes ("MMT", "MMSCM").
    """
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    chunks: list[Chunk] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            if not any(c is not None and str(c).strip() for c in row):
                continue
            rows.append([
                _clean(_fmt_cell(c)) for c in row
            ])
        if not rows:
            continue
        descriptor = _sheet_descriptor(path.stem, sheet_name, rows)
        md = _table_to_markdown(rows)
        body = f"Sheet: {sheet_name}\n{descriptor}\n\n{md}" if descriptor else f"Sheet: {sheet_name}\n{md}"
        chunks.append(Chunk(
            text=body,
            metadata={"section": f"sheet:{sheet_name}"},
        ))
    return chunks


def _sheet_descriptor(filename_stem: str, sheet_name: str, rows: list[list[str]]) -> str:
    """Build a short natural-language line listing the column headers + the
    year range, so retrieval can match "crude oil production over the years"
    against a sheet whose body is just numbers."""
    # Pull plausible header cells from the first few rows.
    head_terms: list[str] = []
    for r in rows[:3]:
        for cell in r:
            v = cell.strip()
            if not v or v in head_terms:
                continue
            # Skip cells that look like a year, percentage, or pure number
            if re.fullmatch(r"\d{4}(-\d{2,4})?", v):
                continue
            if re.fullmatch(r"-?\d+(\.\d+)?%?", v):
                continue
            head_terms.append(v)
        if len(head_terms) > 12:
            break

    # Pull year-like tokens from the leftmost column
    years: list[str] = []
    for r in rows[1:]:
        if not r:
            continue
        first = r[0].strip()
        if re.fullmatch(r"\d{4}(-\d{2,4})?", first) and first not in years:
            years.append(first)

    bits = [f"Workbook: {filename_stem}", f"Sheet: {sheet_name}"]
    if head_terms:
        bits.append("Columns / metrics: " + "; ".join(head_terms[:14]))
    if years:
        bits.append(f"Year range: {years[0]} to {years[-1]}")
    return ". ".join(bits) + "."


def _fmt_cell(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        # Percentage-looking small floats get rounded; otherwise trim trailing zeros
        if abs(value) < 1 and value != 0:
            return f"{value:.4f}".rstrip("0").rstrip(".")
        return f"{value:.4f}".rstrip("0").rstrip(".")
    return str(value)


# ---------- .pdf ----------

def extract_pdf(path: Path) -> list[Chunk]:
    """One chunk per page. Many of the PQ-reply PDFs are scanned — pages
    return empty text and we just skip them (the parallel .docx carries the
    content)."""
    chunks: list[Chunk] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = _clean(page.extract_text() or "")
                if not text or len(text) < 20:
                    continue
                chunks.append(Chunk(
                    text=text,
                    metadata={"section": f"page_{i + 1}"},
                ))
    except Exception as e:
        # Corrupt or unreadable PDFs shouldn't kill the whole ingest
        print(f"[extract_pdf] skipped {path.name}: {e}")
    return chunks


# ---------- .json (synthetic data feeds) ----------

def extract_json(path: Path) -> list[Chunk]:
    """Render a synthetic-data JSON as a markdown table chunk.

    The HSE / Procurement / Workforce JSONs are stable shapes (top-level
    'events' / 'bids' / 'by_function' lists). We dump each list as a
    markdown table so retrieval can semantic-search it and the LLM can
    quote the rows.
    """
    import json

    chunks: list[Chunk] = []
    try:
        data = json.loads(path.read_text())
    except Exception as e:
        print(f"[extract_json] {path}: {e}")
        return []

    def _render_list_of_dicts(name: str, items: list[dict]) -> str:
        if not items:
            return ""
        # Union of keys, preserving order of first occurrence.
        keys: list[str] = []
        for it in items:
            for k in it.keys():
                if k not in keys:
                    keys.append(k)
        header = "| " + " | ".join(keys) + " |"
        sep = "| " + " | ".join(["---"] * len(keys)) + " |"
        rows = []
        for it in items:
            cells = []
            for k in keys:
                v = it.get(k, "")
                if isinstance(v, (list, dict)):
                    v = json.dumps(v, ensure_ascii=False)
                cells.append(_clean(str(v)))
            rows.append("| " + " | ".join(cells) + " |")
        return f"{name}\n\n{header}\n{sep}\n" + "\n".join(rows)

    # Top-level: a few well-known list keys plus a generic dump.
    # Each list becomes its own chunk.
    if isinstance(data, dict):
        for key, val in data.items():
            if isinstance(val, list) and val and isinstance(val[0], dict):
                md = _render_list_of_dicts(f"{key}:", val)
                if md:
                    chunks.append(Chunk(text=md, metadata={"section": key}))
            elif isinstance(val, dict):
                md = "\n".join(f"- **{k}**: {json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else v}"
                                for k, v in val.items())
                chunks.append(Chunk(text=f"{key}:\n\n{md}", metadata={"section": key}))
            elif isinstance(val, (str, int, float, bool)):
                chunks.append(Chunk(text=f"**{key}**: {val}",
                                    metadata={"section": key}))
    elif isinstance(data, list) and data and isinstance(data[0], dict):
        md = _render_list_of_dicts(path.stem, data)
        if md:
            chunks.append(Chunk(text=md, metadata={"section": path.stem}))

    return chunks


# ---------- dispatch ----------

EXTENSIONS = {
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".pdf":  extract_pdf,
    ".json": extract_json,
}


def iter_documents(root: Path) -> Iterable[Path]:
    """Yield every supported document under `root`, skipping temp/lock files."""
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith("~") or path.name.startswith("."):
            continue
        if path.suffix.lower() not in EXTENSIONS:
            continue
        yield path


def extract(path: Path) -> list[Chunk]:
    """Extract chunks for a single file. Returns [] on unsupported / failed."""
    fn = EXTENSIONS.get(path.suffix.lower())
    if not fn:
        return []
    try:
        return fn(path)
    except Exception as e:
        print(f"[extract] failed {path}: {e}")
        return []
